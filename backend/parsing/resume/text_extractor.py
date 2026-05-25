"""Resume file → plain text. PDF via pymupdf; DOCX via python-docx."""

from __future__ import annotations

import io
import logging

logger = logging.getLogger(__name__)


class UnsupportedResumeError(ValueError):
    """Raised when the uploaded file isn't a supported resume format."""


_PDF_EXT = ".pdf"
_DOCX_EXT = ".docx"
_PDF_MIMES = {"application/pdf", "application/x-pdf"}
_DOCX_MIMES = {
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "application/msword",
}


def _detect_kind(filename: str | None, content_type: str | None) -> str:
    """
    Returns "pdf" or "docx". Raises UnsupportedResumeError otherwise.
    Uses extension first (most reliable on Windows), MIME as fallback.
    """
    name = (filename or "").lower().strip()
    if name.endswith(_PDF_EXT):
        return "pdf"
    if name.endswith(_DOCX_EXT):
        return "docx"

    mime = (content_type or "").lower().split(";", 1)[0].strip()
    if mime in _PDF_MIMES:
        return "pdf"
    if mime in _DOCX_MIMES:
        return "docx"

    raise UnsupportedResumeError(
        "Unsupported file type. Upload a PDF or DOCX resume."
    )


def _extract_pdf(data: bytes) -> str:
    # Use the canonical `pymupdf` import name. The legacy `fitz` alias still
    # works when only PyMuPDF is installed, but `fitz` is also the name of an
    # unrelated, broken squatter package on PyPI (`fitz-0.0.1.dev2`) that
    # overwrites PyMuPDF's files when accidentally installed and dies inside
    # its `from frontend import *`. Importing as `pymupdf` sidesteps that.
    import pymupdf

    text_chunks: list[str] = []
    with pymupdf.open(stream=data, filetype="pdf") as doc:
        for page in doc:
            page_text = page.get_text("text") or ""
            if page_text.strip():
                text_chunks.append(page_text)
    return "\n".join(text_chunks).strip()


def _extract_docx(data: bytes) -> str:
    from docx import Document  # python-docx

    doc = Document(io.BytesIO(data))

    parts: list[str] = []
    for paragraph in doc.paragraphs:
        text = (paragraph.text or "").strip()
        if text:
            parts.append(text)
    # python-docx doesn't include table cells in iter_paragraphs, so add them.
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                cell_text = (cell.text or "").strip()
                if cell_text:
                    parts.append(cell_text)
    return "\n".join(parts).strip()


def extract_resume_text(
    *,
    data: bytes,
    filename: str | None,
    content_type: str | None,
) -> str:
    """
    Returns the plain-text contents of the resume. Raises UnsupportedResumeError
    for unsupported formats and ValueError if the file is empty / unreadable.
    """
    if not data:
        raise ValueError("Uploaded file is empty.")

    kind = _detect_kind(filename, content_type)
    try:
        text = _extract_pdf(data) if kind == "pdf" else _extract_docx(data)
    except Exception as exc:
        logger.exception("resume text extraction failed (%s)", kind)
        raise ValueError(f"Could not read {kind.upper()} file: {exc}") from exc

    if not text.strip():
        raise ValueError(
            "No readable text found in the resume. "
            "If your PDF is a scan, export it as a text PDF first."
        )
    return text
